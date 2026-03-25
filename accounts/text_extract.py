"""Extraction de texte depuis PDF et Word (.docx) pour l'analyse de plagiat."""

from __future__ import annotations


def extract_text_from_upload(uploaded_file) -> str:
    """
    Extrait le texte d'un fichier uploadé.
    Lève ValueError avec un message utilisateur si le format n'est pas géré.
    """
    name = (uploaded_file.name or "").lower()

    if name.endswith(".pdf"):
        return _extract_pdf(uploaded_file)
    if name.endswith(".docx"):
        return _extract_docx(uploaded_file)
    if name.endswith(".doc"):
        raise ValueError(
            "Les fichiers .doc (Word 97–2003) ne sont pas pris en charge. "
            "Enregistrez votre document au format .docx ou exportez-le en PDF."
        )
    raise ValueError("Format de fichier non reconnu. Utilisez un PDF ou un fichier Word (.docx).")


def _extract_pdf(file_obj) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_obj)
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _extract_docx(file_obj) -> str:
    from docx import Document

    doc = Document(file_obj)
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)
